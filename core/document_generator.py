import os
import random
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Pt
from reportlab.platypus.tables import ParagraphStyle
from constants import DOCUMENT_FONTS, MIN_FONT_SIZE, MAX_FONT_SIZE

from utils.text_utils import convert_markdown_bold_to_html, split_markdown_bold


class DocumentGenerator:
    """Generates PDF and Word documents with answers."""
    
    @staticmethod
    def generate_pdf(
        exercises_with_answers: list[tuple[str, str]],
        student_name: str,
        group: str,
        output_filename: str
    ) -> str:
        """
        Generate PDF file with answers.
        
        Args:
            exercises_with_answers: List of (exercise_text, answer_text) tuples
            student_name: Student name
            group: Group/class name
            output_filename: Output filename (without extension)
            
        Returns:
            Path to generated PDF file
        """
        if not output_filename.endswith('.pdf'):
            output_filename += '.pdf'
        
        output_path = os.path.join(os.getcwd(), output_filename)
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        
        # Add header information
        story.append(Paragraph(f"<b>Student:</b> {student_name}", normal_style))
        story.append(Spacer(1, 0.1*inch))
        story.append(Paragraph(f"<b>Group:</b> {group}", normal_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Add exercises and answers
        for exercise_text, answer_text in exercises_with_answers:
            formatted_answer = convert_markdown_bold_to_html(answer_text)
            formatted_answer = formatted_answer.replace('\n', '<br/>')

            font_name = random.choice(DOCUMENT_FONTS) 

            font_name = font_name if font_name != "Times New Roman" else "Times-Roman"

            paragraph_style = ParagraphStyle(
                        name='RandomStyle',
                        fontName=font_name,
                        fontSize=random.randint(MIN_FONT_SIZE, MAX_FONT_SIZE),
                        leading=14
                    )
            
            story.append(Paragraph(formatted_answer, paragraph_style))
            story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
        
        return output_path
    
    @staticmethod
    def generate_word(
        exercises_with_answers: list[tuple[str, str]],
        student_name: str,
        group: str,
        output_filename: str
    ) -> str:
        """
        Generate Word document file with answers.
        
        Args:
            exercises_with_answers: List of (exercise_text, answer_text) tuples
            student_name: Student name
            group: Group/class name
            output_filename: Output filename (without extension)
            
        Returns:
            Path to generated Word file
        """
        if not output_filename.endswith('.docx'):
            output_filename += '.docx'
        
        output_path = os.path.join(os.getcwd(), output_filename)
        
        # Create Word document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = random.choice(DOCUMENT_FONTS)
        font.size = Pt(random.randint(MIN_FONT_SIZE, MAX_FONT_SIZE))
        
        # Add header information
        p = doc.add_paragraph()
        p.add_run('Student: ').bold = True
        p.add_run(student_name)
        
        p = doc.add_paragraph()
        p.add_run('Group: ').bold = True
        p.add_run(group)
        
        doc.add_paragraph()
        
        # Add exercises and answers
        for idx, (exercise_text, answer_text) in enumerate(exercises_with_answers, 1):
            answer_lines = answer_text.split('\n')
            for line in answer_lines:
                line = line.strip()
                if line:
                    p = doc.add_paragraph()
                    parts = split_markdown_bold(line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        elif part:
                            p.add_run(part)
                else:
                    doc.add_paragraph()
            
            if idx < len(exercises_with_answers):
                doc.add_paragraph()
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    @staticmethod
    def generate_output_filename(student_name: str, group: str, custom_filename: str = "") -> str:
        """
        Generate output filename from student name and group.
        
        Args:
            student_name: Student name
            group: Group/class name
            custom_filename: Optional custom filename
            
        Returns:
            Generated filename (without extension)
        """
        if custom_filename:
            return custom_filename
        
        name_clean = student_name.replace(' ', '_')
        group_clean = group.replace(' ', '_')
        return f"{name_clean}_{group_clean}"

